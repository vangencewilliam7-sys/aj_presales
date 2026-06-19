import os
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_doc(title, content_blocks, filename):
    doc = Document()
    
    # Title
    heading = doc.add_heading(title, 0)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    for block in content_blocks:
        if block['type'] == 'h1':
            doc.add_heading(block['text'], level=1)
        elif block['type'] == 'h2':
            doc.add_heading(block['text'], level=2)
        elif block['type'] == 'h3':
            doc.add_heading(block['text'], level=3)
        elif block['type'] == 'p':
            doc.add_paragraph(block['text'])
        elif block['type'] == 'code':
            p = doc.add_paragraph()
            run = p.add_run(block['text'])
            run.font.name = 'Courier New'
            run.font.size = Pt(9)
        elif block['type'] == 'bullet':
            doc.add_paragraph(block['text'], style='List Bullet')
            
    doc.save(filename)

# --- Document 1: RAG Ingestion Pipeline ---
doc1_blocks = [
    {'type': 'h1', 'text': 'High-Level Design (HLD)'},
    {'type': 'p', 'text': 'The RAG Ingestion Pipeline is responsible for processing heterogeneous knowledge sources (YouTube transcripts and PDF books) into a structured, easily retrievable format. It utilizes a Two-Track Chunking Engine to handle timestamped data differently from continuous text, ensuring optimal context window usage.'},
    {'type': 'h2', 'text': 'Architecture Diagram (Flow)'},
    {'type': 'code', 'text': '''
Raw Data (TXT/DOCX/PDF) 
  --> Parser (Extract Text & Metadata)
    --> Two-Track Chunking Engine:
        Track A: Timestamp-based (Video)
        Track B: Character/Page-based (Book)
          --> OpenAI text-embedding-3-small
            --> Supabase (knowledge_sources & knowledge_chunks)
    '''},
    {'type': 'h1', 'text': 'Low-Level Design (LLD)'},
    {'type': 'h2', 'text': '1. Parsing Layer'},
    {'type': 'bullet', 'text': 'Video Transcripts: Parsed using standard regex to separate headers from transcript text.'},
    {'type': 'bullet', 'text': 'PDF Books: Extracted page-by-page using PyMuPDF (fitz).'},
    {'type': 'h2', 'text': '2. Two-Track Chunking Engine'},
    {'type': 'bullet', 'text': 'Track A (Timestamped): Splits text using [HH:MM:SS] regex markers. Groups segments until ~800 characters are reached. Location marker is the starting timestamp.'},
    {'type': 'bullet', 'text': 'Track B (Character/Page): Recursive splitting with overlap (800 chars, 200 overlap) ensuring word boundaries are respected. Location marker is the page or segment number.'},
    {'type': 'h2', 'text': '3. Embedding & Storage'},
    {'type': 'bullet', 'text': 'OpenAI Embeddings: Uses text-embedding-3-small for 1536-dimensional dense vectors.'},
    {'type': 'bullet', 'text': 'Supabase Batch Insertion: Inserts parent metadata into knowledge_sources, then batches chunks (20 at a time) into knowledge_chunks.'},
    {'type': 'h1', 'text': 'Code References'},
    {'type': 'p', 'text': 'backend/ingest_data.py'},
    {'type': 'code', 'text': 'def chunk_timestamped(transcript_text, chunk_size=800)\ndef chunk_by_characters(transcript_text, chunk_size=800, overlap=100)'}
]

# --- Document 2: Hybrid RAG Retrieval System ---
doc2_blocks = [
    {'type': 'h1', 'text': 'High-Level Design (HLD)'},
    {'type': 'p', 'text': 'The Hybrid RAG Retrieval System ensures zero-trust grounding by fetching relevant chunks from the database. It combines semantic search (dense vectors) with exact keyword fallback to ensure no specific terminology is missed during context retrieval.'},
    {'type': 'h2', 'text': 'Architecture Diagram (Flow)'},
    {'type': 'code', 'text': '''
Expert Answer 
  --> OpenAI Embeddings (text-embedding-3-small)
    --> Supabase RPC `match_knowledge_chunks` (Cosine Similarity)
      --> If results empty:
          --> Fallback: Supabase `ilike` (Keyword Search)
        --> Return Context Text & Chunk Metadata
    '''},
    {'type': 'h1', 'text': 'Low-Level Design (LLD)'},
    {'type': 'h2', 'text': '1. Vector Search (Primary)'},
    {'type': 'bullet', 'text': 'Uses pgvector HNSW index for fast nearest-neighbor search.'},
    {'type': 'bullet', 'text': 'RPC Function: match_knowledge_chunks takes query_embedding, match_threshold (0.3), and match_count (4).'},
    {'type': 'h2', 'text': '2. Keyword Fallback (Secondary)'},
    {'type': 'bullet', 'text': 'If the vector search yields 0 results, it falls back to a SQL ILIKE search on the first 15 characters of the query.'},
    {'type': 'h2', 'text': '3. Context Assembly'},
    {'type': 'bullet', 'text': 'Joins chunk content with parent source title and location markers to provide fully cited context to the LLM.'},
    {'type': 'h1', 'text': 'Code References'},
    {'type': 'p', 'text': 'backend/app.py'},
    {'type': 'code', 'text': 'def hybrid_rag_fetch(query: str, top_k: int = 4) -> dict:'}
]

# --- Document 3: Hierarchical Knowledge Schema ---
doc3_blocks = [
    {'type': 'h1', 'text': 'High-Level Design (HLD)'},
    {'type': 'p', 'text': 'A relational, parent-child schema in PostgreSQL (via Supabase) utilizing pgvector for dense embeddings and tsvector for sparse text search. It ensures data normalization and efficient hierarchical querying.'},
    {'type': 'h2', 'text': 'Entity Relationship Diagram'},
    {'type': 'code', 'text': '''
[knowledge_sources] (Parent)
- id (UUID, PK)
- source_type (youtube/book)
- title, author, global_summary
      |
     1:N
      |
[knowledge_chunks] (Child)
- id (UUID, PK)
- source_id (UUID, FK)
- content (TEXT)
- location_marker (VARCHAR)
- embedding (VECTOR 1536)
- fts_tokens (TSVECTOR)
    '''},
    {'type': 'h1', 'text': 'Low-Level Design (LLD)'},
    {'type': 'bullet', 'text': 'Parent Table: knowledge_sources holds top-level metadata. Allows filtering by source_type before vector search.'},
    {'type': 'bullet', 'text': 'Child Table: knowledge_chunks stores the actual text chunks. The fts_tokens column is generated automatically using PostgreSQL to_tsvector().'},
    {'type': 'bullet', 'text': 'Indexes: HNSW index on the embedding column using vector_cosine_ops for fast similarity search. GIN index on fts_tokens for rapid full-text search.'},
    {'type': 'h1', 'text': 'Code References'},
    {'type': 'p', 'text': 'docs/ingestion.md (SQL DDL Statements)'}
]

# --- Document 4: Pre-Interview Research Scan ---
doc4_blocks = [
    {'type': 'h1', 'text': 'High-Level Design (HLD)'},
    {'type': 'p', 'text': 'Before an interview, the AI Journalist simulates reading a briefing by sampling chunks from the knowledge base. This ensures the script is grounded in the actual database rather than LLM hallucinations.'},
    {'type': 'h2', 'text': 'Architecture Diagram (Flow)'},
    {'type': 'code', 'text': '''
Trigger: /prepare-interview
  --> Fetch all 9 knowledge sources
    --> For each source:
        --> Fetch chunks at indices [0, len//2, len-1] (Start, Middle, End)
          --> Compile 27-chunk Research Briefing
            --> Feed to Theme Extraction LLM
    '''},
    {'type': 'h1', 'text': 'Low-Level Design (LLD)'},
    {'type': 'bullet', 'text': 'Sampling Logic: Instead of feeding the entire DB (which exceeds context limits), it samples the first, middle, and last chunks of every source.'},
    {'type': 'bullet', 'text': 'Briefing Assembly: Combines source_title, source_type, location, and the first 400 characters of the content.'},
    {'type': 'h1', 'text': 'Code References'},
    {'type': 'p', 'text': 'backend/app.py'},
    {'type': 'code', 'text': 'async def research_scan() -> dict:'}
]

# --- Document 5: Script Crafting Engine ---
doc5_blocks = [
    {'type': 'h1', 'text': 'High-Level Design (HLD)'},
    {'type': 'p', 'text': 'The Script Crafting Engine generates a 30-35 question blueprint divided into 4 psychological phases. It mandates chunk attribution, ensuring every scripted question is tied to a specific piece of source evidence.'},
    {'type': 'h2', 'text': 'Architecture Diagram (Flow)'},
    {'type': 'code', 'text': '''
Extracted Themes + Research Briefing
  --> SCRIPT_CRAFTING_PROMPT
    --> GPT-4o-mini
      --> Outputs 4 Phases (JSON)
          - Phase 1: Warmup
          - Phase 2: Deep Dives
          - Phase 3: Challenge
          - Phase 4: Synthesis
            --> Insert into `interview_scripts` DB table
    '''},
    {'type': 'h1', 'text': 'Low-Level Design (LLD)'},
    {'type': 'bullet', 'text': 'Structured Output: Enforces a strict nested JSON schema detailing the phase goals, emotional triggers, contingencies, and estimated duration.'},
    {'type': 'bullet', 'text': 'Chunk Attribution: The LLM must output a `chunk_attribution` object for each question, including `why_this_chunk` to simulate editorial reasoning.'},
    {'type': 'bullet', 'text': 'Archetype Integration: Applies different interviewing styles (e.g., Lex Fridman, Dwarkesh Patel) based on the prompt directives.'},
    {'type': 'h1', 'text': 'Code References'},
    {'type': 'p', 'text': 'backend/prompts.py & backend/app.py (/prepare-interview)'},
    {'type': 'code', 'text': 'SCRIPT_CRAFTING_PROMPT'}
]

# --- Document 6: Intent Classification System ---
doc6_blocks = [
    {'type': 'h1', 'text': 'High-Level Design (HLD)'},
    {'type': 'p', 'text': "A fast-path routing mechanism that evaluates an expert's response before full evaluation. It determines if the expert wants to skip the question or provide a substantive answer."},
    {'type': 'h2', 'text': 'Architecture Diagram (Flow)'},
    {'type': 'code', 'text': '''
Expert Answer 
  --> INTENT_CLASSIFIER_PROMPT (GPT-4o-mini, max_tokens=30)
    --> JSON: {"intent": "skip" | "substantive"}
      --> If "skip":
          Increment completed_count, Move to next scripted question
      --> If "substantive":
          Proceed to Full Evaluation (RAG + SCRIPT_AWARE_EVALUATION)
    '''},
    {'type': 'h1', 'text': 'Low-Level Design (LLD)'},
    {'type': 'bullet', 'text': 'Speed Optimization: Uses a low temperature (0.0) and small max_tokens limit to ensure sub-second classification latency.'},
    {'type': 'bullet', 'text': 'Skip Logic: If intent is "skip", it updates the script progress in the DB and immediately issues the next scripted question, saving a heavy RAG query and full evaluation cycle.'},
    {'type': 'h1', 'text': 'Code References'},
    {'type': 'p', 'text': 'backend/app.py'},
    {'type': 'code', 'text': 'intent_res = llm_fast.invoke(INTENT_CLASSIFIER_PROMPT...)'}
]

# --- Document 7: Interviewer Persona System ---
doc7_blocks = [
    {'type': 'h1', 'text': 'High-Level Design (HLD)'},
    {'type': 'p', 'text': 'The persona system dictates the cognitive and conversational style of the AI. It uses the "Tacit Knowledge Protocol" to bypass rehearsed answers by anchoring questions in emotion.'},
    {'type': 'h2', 'text': 'Architecture Diagram (Flow)'},
    {'type': 'code', 'text': '''
System Base: JOURNALIST_BASE_PERSONA
  + Domain: Enterprise Solutions Architecture
  + Strategies:
    - Lex Fridman (Human pressure, silence)
    - Dwarkesh Patel (Contrasting philosophies)
    - O'Shaughnessy (SOP/Framework extraction)
    - Shane Parrish (Mental models)
      --> Applied in GENERATION_PHASE_PROMPT
    '''},
    {'type': 'h1', 'text': 'Low-Level Design (LLD)'},
    {'type': 'bullet', 'text': 'Tacit Knowledge Protocol: Forces the LLM to avoid "How do you handle X?" and instead ask "Walk me through the moment when X failed."'},
    {'type': 'bullet', 'text': 'Zero-Trust Grounding: Explicit instructions to never hallucinate and heavily rely on the `expert_answer` and `db_context`.'},
    {'type': 'h1', 'text': 'Code References'},
    {'type': 'p', 'text': 'backend/prompts.py'},
    {'type': 'code', 'text': 'JOURNALIST_BASE_PERSONA\nSTRATEGY_RULES'}
]

# --- Document 8: Decision Transparency Layer ---
doc8_blocks = [
    {'type': 'h1', 'text': 'High-Level Design (HLD)'},
    {'type': 'p', 'text': "A frontend feature (Glass-Box UI) that exposes the agentic thought process to the user. It demystifies the AI's choices by rendering its internal monologue, tangent detection, and RAG contexts directly below the chat bubble."},
    {'type': 'h2', 'text': 'Architecture Diagram (Flow)'},
    {'type': 'code', 'text': '''
Backend `/generate-question` Response:
{ "question": "...", "decision": { "internal_monologue": "...", "action": "...", "chunks_used": [...] } }
  --> React State (messages array)
    --> Render Message Bubble
      --> "Decision Log" Toggle Button
        --> Expands to show metadata (Monologue, RAG citations, Next Action)
    '''},
    {'type': 'h1', 'text': 'Low-Level Design (LLD)'},
    {'type': 'bullet', 'text': 'React State: Tracks `openDecisionId` to allow expanding/collapsing individual decision logs without cluttering the UI.'},
    {'type': 'bullet', 'text': 'Metadata Payload: The backend passes the entire JSON output of the Evaluation prompt alongside the generated text.'},
    {'type': 'bullet', 'text': 'Styling: Implemented as a dark, monospace-styled card to visually separate it from conversational flow.'},
    {'type': 'h1', 'text': 'Code References'},
    {'type': 'p', 'text': 'frontend/src/App.tsx'},
    {'type': 'code', 'text': '<button className="decision-toggle" ...>\n{msg.decision.internal_monologue}'}
]

import os
os.makedirs('docs/Architecture_Docs', exist_ok=True)

create_doc('1. RAG Ingestion Pipeline', doc1_blocks, 'docs/Architecture_Docs/1_RAG_Ingestion_Pipeline.docx')
create_doc('2. Hybrid RAG Retrieval System', doc2_blocks, 'docs/Architecture_Docs/2_Hybrid_RAG_Retrieval.docx')
create_doc('3. Hierarchical Knowledge Schema', doc3_blocks, 'docs/Architecture_Docs/3_Hierarchical_Knowledge_Schema.docx')
create_doc('4. Pre-Interview Research Scan', doc4_blocks, 'docs/Architecture_Docs/4_Pre_Interview_Research_Scan.docx')
create_doc('5. Script Crafting Engine', doc5_blocks, 'docs/Architecture_Docs/5_Script_Crafting_Engine.docx')
create_doc('6. Intent Classification System', doc6_blocks, 'docs/Architecture_Docs/6_Intent_Classification_System.docx')
create_doc('7. Interviewer Persona System', doc7_blocks, 'docs/Architecture_Docs/7_Interviewer_Persona_System.docx')
create_doc('8. Decision Transparency Layer', doc8_blocks, 'docs/Architecture_Docs/8_Decision_Transparency_Layer.docx')

print("Generated 8 Word documents in docs/Architecture_Docs/")
